"""
勤勉 AI — Flask 应用入口 v4
============================
新增功能：
  - 知识库长期记忆（KnowledgeBase）
  - 对话持久化存储（ConversationStore）
  - 知识库/大模型独立开关
  - 对话历史管理（创建、查询、删除）
  - 全屏 AI 助手页面（赛博朋克风格）
"""

from __future__ import annotations

import base64
import binascii
import datetime
import functools
import io
import json
import mimetypes
import os
import re
import subprocess
import tempfile
import threading
import time
from pathlib import Path
from typing import Any
from urllib.parse import urlsplit

import flask
from flask import Flask, Response, g, jsonify, request, send_file, session, stream_with_context
from flask_cors import CORS

from qinmian.agent import QinmianAgent
from qinmian.analytics import (
    ConflictResolver,
    CourseDifficultyDB,
    CourseHardnessAnalyzer,
    CreditChecker,
    ProfessorMatcher,
)
from qinmian.auth_store import (
    UserStore,
    delete_user_data,
    load_or_create_session_secret,
    user_data_path,
)
from qinmian.conversation_store import (
    add_message,
    create_conversation,
    delete_conversation,
    delete_message,
    get_conversation,
    get_or_create_active,
    list_conversations,
    rename_conversation,
)
from qinmian.data_store import PROJECT_ROOT, QinmianDataStore
from qinmian.export_service import build_export
from qinmian.knowledge_base import KnowledgeBase
from qinmian.personas import public_personas
from qinmian.persistence import database_enabled, delete_owner_documents
from qinmian.planner import CareerPlanner
from qinmian.tools import FunctionCallExecutor, get_function_schemas
from qinmian.user_llm import UserLLMConfigStore
from qinmian.user_runtime import UserRuntimeStoreManager

# ── Flask 应用 ───────────────────────────────────────────────────────
app = Flask(__name__, static_folder=str(PROJECT_ROOT / "static"), static_url_path="")
app.config.update(
    SECRET_KEY=load_or_create_session_secret(),
    PERMANENT_SESSION_LIFETIME=datetime.timedelta(days=14),
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=os.getenv("QINMIAN_COOKIE_SECURE", "0").lower()
    in {"1", "true", "yes", "on"},
)

cors_origins = [
    origin.strip()
    for origin in os.getenv("QINMIAN_CORS_ORIGINS", "").split(",")
    if origin.strip()
]
if cors_origins:
    CORS(
        app,
        resources={r"/api/*": {"origins": cors_origins}},
        supports_credentials=True,
    )

# ── 全局单例 ─────────────────────────────────────────────────────────
STORE = QinmianDataStore()
AGENT = QinmianAgent(STORE)
USER_STORE = UserStore()
_KNOWLEDGE_BASES: dict[str, KnowledgeBase] = {}
_KNOWLEDGE_LOCK = threading.RLock()
COURSE_ANALYZER = CourseHardnessAnalyzer(STORE)
COURSE_DIFFICULTY = CourseDifficultyDB()
PROF_MATCHER = ProfessorMatcher(STORE)
CREDIT_CHECKER = CreditChecker(STORE)
CONFLICT_RESOLVER = ConflictResolver(STORE)
CAREER_PLANNER = CareerPlanner(STORE)
FC_EXECUTOR = FunctionCallExecutor(STORE)
USER_RUNTIME_STORES = UserRuntimeStoreManager(STORE)
USER_LLM_CONFIGS = UserLLMConfigStore(str(app.config["SECRET_KEY"]))

DATA_RELEASE_DATE = "2026-07-23"
DATA_DISCLAIMER = (
    "课程、职业与教师分析仅供学业规划参考；培养方案、开课安排和教师信息"
    "请以华侨大学教务处及学院最新通知为准。"
)
SIMULATION_FEATURES = {
    "career_plan": "规则与数据驱动的职业规划参考，不代表就业承诺",
    "seat_monitor": "演示用模拟余量，不连接学校实时选课系统",
    "conflict_resolution": "决策辅助建议，不会直接修改教务系统课表",
    "generated_curriculum": "部分课程按培养模板补全，需以官方方案复核",
}
_RATE_LIMIT_LOCK = threading.RLock()
_RATE_LIMIT_BUCKETS: dict[str, list[float]] = {}


def _client_address() -> str:
    forwarded = request.headers.get("X-Forwarded-For", "")
    return (forwarded.split(",", 1)[0].strip() or request.remote_addr or "unknown")[:128]


def _rate_limit(key: str, limit: int, window_seconds: int):
    """Return a 429 response when a small in-process request budget is exhausted."""
    if app.config.get("TESTING"):
        return None
    now = time.time()
    cutoff = now - window_seconds
    with _RATE_LIMIT_LOCK:
        bucket = [stamp for stamp in _RATE_LIMIT_BUCKETS.get(key, []) if stamp > cutoff]
        if len(bucket) >= limit:
            retry_after = max(1, int(bucket[0] + window_seconds - now) + 1)
            response = jsonify({
                "error": f"操作过于频繁，请在 {retry_after} 秒后重试",
                "code": "rate_limited",
                "retry_after": retry_after,
            })
            response.status_code = 429
            response.headers["Retry-After"] = str(retry_after)
            return response
        bucket.append(now)
        _RATE_LIMIT_BUCKETS[key] = bucket
    return None


def _clear_rate_limit(key: str) -> None:
    with _RATE_LIMIT_LOCK:
        _RATE_LIMIT_BUCKETS.pop(key, None)


def _data_governance() -> dict[str, Any]:
    return {
        "release_date": DATA_RELEASE_DATE,
        "disclaimer": DATA_DISCLAIMER,
        "official_system_connected": False,
        "simulation_features": SIMULATION_FEATURES,
        "source_levels": {
            "catalog": "公开目录与导入数据",
            "credits": "官方表格匹配时优先使用，否则标记为模板",
            "curriculum": "官方字段与规划模板混合，生成内容会明确标记",
        },
    }


# ═════════════════════════════════════════════════════════════════════
# 工具函数
# ═════════════════════════════════════════════════════════════════════

def _query_arg(key: str, default: str = "") -> str:
    return request.args.get(key, default)


def _query_int(key: str, default: int = 0) -> int:
    try:
        return int(request.args.get(key, str(default)))
    except (ValueError, TypeError):
        return default


def _body_json() -> dict[str, Any]:
    return request.get_json(silent=True) or {}


def _langchain_available() -> bool:
    try:
        import langchain  # noqa
        return True
    except ImportError:
        return False


def _current_user_id() -> str:
    user = getattr(g, "current_user", None)
    if not user:
        raise RuntimeError("authentication required")
    return str(user["id"])


def _knowledge_base_for(user_id: str | None = None) -> KnowledgeBase:
    owner_id = user_id or _current_user_id()
    with _KNOWLEDGE_LOCK:
        knowledge_base = _KNOWLEDGE_BASES.get(owner_id)
        if knowledge_base is None:
            store_path = user_data_path(owner_id) / "knowledge_base" / "records.json"
            knowledge_base = KnowledgeBase(store_path, owner_id=owner_id)
            knowledge_base.index_major_catalog(STORE)
            _KNOWLEDGE_BASES[owner_id] = knowledge_base
        return knowledge_base


def _llm_client_for_user(user_id: str | None = None):
    owner_id = user_id or _current_user_id()
    try:
        return USER_LLM_CONFIGS.get_client(owner_id, AGENT.llm)
    except ValueError as exc:
        from qinmian.llm import LLMClient

        client = LLMClient({
            "provider": "personal",
            "base_url": "https://api.openai.com/v1",
            "model": "unconfigured",
            "vision_model": "unconfigured",
            "display_name": "个人 API 配置不可用",
            "api_key": "",
        })
        client.last_error = str(exc)
        return client


def _llm_status_for_user() -> dict[str, Any]:
    user_id = _current_user_id()
    client = _llm_client_for_user(user_id)
    settings = USER_LLM_CONFIGS.public_settings(user_id)
    status = dict(client.status())
    configured = bool(status.get("enabled"))
    user_enabled = bool(session.get("llm_enabled", True))
    status["configured"] = configured
    status["user_enabled"] = user_enabled
    status["enabled"] = configured and user_enabled
    status["source"] = settings["source"]
    status["has_personal_config"] = settings["source"] == "personal"
    return status


@app.before_request
def load_authenticated_user():
    user_id = str(session.get("user_id", ""))
    g.current_user = USER_STORE.get(user_id) if user_id else None
    if user_id and not g.current_user:
        session.clear()

    if not request.path.startswith("/api/"):
        return None
    if request.method not in {"GET", "HEAD", "OPTIONS"}:
        origin = request.headers.get("Origin", "").rstrip("/")
        fetch_site = request.headers.get("Sec-Fetch-Site", "")
        origin_host = urlsplit(origin).netloc if origin else ""
        if (origin_host and origin_host != request.host) or fetch_site == "cross-site":
            return jsonify({
                "error": "已拒绝跨站请求",
                "code": "cross_site_request_blocked",
            }), 403
    public_auth_paths = {
        "/api/auth/me",
        "/api/auth/register",
        "/api/auth/login",
    }
    if request.path in public_auth_paths:
        return None
    if not g.current_user:
        return jsonify({"error": "请先登录", "code": "authentication_required"}), 401
    return None


# ═════════════════════════════════════════════════════════════════════
# 静态文件与前端入口
# ═════════════════════════════════════════════════════════════════════

@app.route("/")
def index():
    return flask.send_from_directory(app.static_folder, "index.html")


@app.route("/health")
def health():
    return jsonify({
        "status": "ok",
        "service": "qinmian-ai-v4",
        "major_count": len(STORE.majors),
        "model_configured": AGENT.llm_status().get("enabled", False),
        "storage": "postgresql" if database_enabled() else "local",
        "data_release_date": DATA_RELEASE_DATE,
    })


# ═════════════════════════════════════════════════════════════════════
# API: 用户认证
# ═════════════════════════════════════════════════════════════════════

def _start_user_session(user: dict[str, Any]) -> None:
    session.clear()
    session.permanent = True
    session["user_id"] = user["id"]
    session["llm_enabled"] = True


@app.route("/api/auth/me")
def api_auth_me():
    if not g.current_user:
        return jsonify({
            "authenticated": False,
            "registration_available": True,
            "user_count": USER_STORE.count(),
        })
    return jsonify({
        "authenticated": True,
        "user": g.current_user,
    })


@app.route("/api/auth/register", methods=["POST"])
def api_auth_register():
    body = _body_json()
    limited = _rate_limit(
        f"register:{_client_address()}",
        limit=5,
        window_seconds=60 * 60,
    )
    if limited:
        return limited
    password = str(body.get("password", ""))
    confirmation = str(body.get("password_confirm", password))
    if password != confirmation:
        return jsonify({"error": "两次输入的密码不一致"}), 400
    try:
        user, _ = USER_STORE.register(body.get("username", ""), password)
    except ValueError as exc:
        message = str(exc)
        status_code = 409 if message == "用户名已存在" else 400
        return jsonify({"error": message}), status_code
    _start_user_session(user)
    return jsonify({
        "status": "ok",
        "user": user,
        "inherited_legacy_data": False,
    }), 201


@app.route("/api/auth/login", methods=["POST"])
def api_auth_login():
    body = _body_json()
    username = str(body.get("username", "")).strip().casefold()
    limit_key = f"login:{_client_address()}:{username}"
    limited = _rate_limit(limit_key, limit=10, window_seconds=10 * 60)
    if limited:
        return limited
    user = USER_STORE.authenticate(username, body.get("password", ""))
    if not user:
        return jsonify({"error": "用户名或密码错误"}), 401
    _clear_rate_limit(limit_key)
    _start_user_session(user)
    return jsonify({"status": "ok", "user": user})


@app.route("/api/auth/logout", methods=["POST"])
def api_auth_logout():
    session.clear()
    return jsonify({"status": "ok"})


@app.route("/api/auth/password", methods=["POST"])
def api_auth_password():
    body = _body_json()
    new_password = str(body.get("new_password", ""))
    confirmation = str(body.get("new_password_confirm", ""))
    if new_password != confirmation:
        return jsonify({"error": "两次输入的新密码不一致"}), 400
    limited = _rate_limit(
        f"password:{_current_user_id()}",
        limit=5,
        window_seconds=10 * 60,
    )
    if limited:
        return limited
    try:
        changed = USER_STORE.change_password(
            _current_user_id(),
            body.get("current_password", ""),
            new_password,
        )
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    if not changed:
        return jsonify({"error": "当前密码错误"}), 401
    return jsonify({"status": "ok", "message": "密码修改成功"})


@app.route("/api/auth/account", methods=["DELETE"])
def api_auth_account():
    body = _body_json()
    user_id = _current_user_id()
    password = body.get("password", "")
    if str(body.get("confirmation", "")).strip() != "注销账号":
        return jsonify({"error": "请输入“注销账号”进行确认"}), 400
    if not USER_STORE.authenticate(g.current_user["username"], password):
        return jsonify({"error": "密码错误，账号未注销"}), 401
    USER_LLM_CONFIGS.clear_settings(user_id)
    USER_RUNTIME_STORES.forget(user_id)
    with _KNOWLEDGE_LOCK:
        _KNOWLEDGE_BASES.pop(user_id, None)
    if database_enabled():
        delete_owner_documents(user_id)
    else:
        delete_user_data(user_id)
    if not USER_STORE.delete_account(user_id, password):
        return jsonify({"error": "账号状态已变化，请重新登录后再试"}), 409
    session.clear()
    return jsonify({
        "status": "ok",
        "message": "账号及个人数据已永久删除",
    })


# ═════════════════════════════════════════════════════════════════════
# API: 元数据
# ═════════════════════════════════════════════════════════════════════

@app.route("/api/meta")
def api_meta():
    knowledge_base = _knowledge_base_for()
    llm_status = _llm_status_for_user()
    return jsonify({
        "name": "勤勉",
        "source": STORE.majors_doc["source"],
        "campuses": STORE.campuses(),
        "colleges": STORE.colleges(),
        "disciplines": STORE.disciplines(),
        "major_count": len(STORE.majors),
        "faculty": STORE.faculty_profiles_doc.get("source", {}),
        "personas": public_personas(),
        "llm": llm_status,
        "knowledge_base": knowledge_base.status(),
        "user": g.current_user,
        "features": {
            "flask": True,
            "sse": True,
            "function_calling": bool(llm_status.get("enabled")),
            "langchain_available": _langchain_available(),
            "knowledge_base": True,
            "conversations": True,
            "exports": ["markdown", "csv", "docx", "pdf", "xls"],
        },
        "data_quality": STORE.data_quality_summary(),
        "data_governance": _data_governance(),
    })


# ═════════════════════════════════════════════════════════════════════
# API: 专业
# ═════════════════════════════════════════════════════════════════════

@app.route("/api/majors")
def api_majors():
    return jsonify(
        STORE.list_majors(
            q=_query_arg("q"),
            campus=_query_arg("campus"),
            college=_query_arg("college"),
            discipline=_query_arg("discipline"),
        )
    )


@app.route("/api/majors/<path:major_id>")
def api_major(major_id: str):
    major = STORE.get_major(major_id)
    if not major:
        return jsonify({"error": "major not found"}), 404
    return jsonify(major)


@app.route("/api/curriculum/<path:major_id>")
def api_curriculum(major_id: str):
    try:
        payload = STORE.curriculum_for(
            major_id,
            _query_arg("student_type", "domestic"),
        )
        payload["evidence"] = {
            "mode": "reference_template",
            "official_schedule": False,
            "release_date": DATA_RELEASE_DATE,
            "notice": SIMULATION_FEATURES["generated_curriculum"],
        }
        return jsonify(payload)
    except KeyError as e:
        return jsonify({"error": str(e)}), 404


# ═════════════════════════════════════════════════════════════════════
# API: 热门方向
# ═════════════════════════════════════════════════════════════════════

@app.route("/api/hot")
def api_hot():
    return jsonify(STORE.hot_directions())


@app.route("/api/careers")
def api_careers():
    roles = STORE.career_roles()
    categories: dict[str, int] = {}
    for role in roles:
        category = role["category"]
        categories[category] = categories.get(category, 0) + 1
    return jsonify({
        "count": len(roles),
        "roles": roles,
        "categories": [
            {"name": name, "count": count}
            for name, count in sorted(categories.items())
        ],
        "notice": "职业画像用于学业规划参考，不代表就业承诺；资格与招聘要求请以最新规定为准。",
    })


@app.route("/api/careers/recommendations")
def api_career_recommendations():
    major_id = request.args.get("major_id", "").strip()
    if not major_id:
        return jsonify({
            "error": "major_id is required",
            "code": "major_id_required",
        }), 400
    try:
        limit = int(request.args.get("limit", "6"))
    except ValueError:
        limit = 6
    try:
        return jsonify(CAREER_PLANNER.recommend_for_major(major_id, limit))
    except KeyError:
        return jsonify({
            "error": "未找到该专业",
            "code": "major_not_found",
        }), 404


# ═════════════════════════════════════════════════════════════════════
# API: 余位监控
# ═════════════════════════════════════════════════════════════════════

@app.route("/api/seats")
def api_seats():
    user_store = USER_RUNTIME_STORES.get(_current_user_id())
    return jsonify({
        "offerings": user_store.offerings(),
        "watchers": user_store.watchers,
        "events": user_store.events[-12:],
        "evidence": {
            "mode": "simulation",
            "official_system_connected": False,
            "notice": SIMULATION_FEATURES["seat_monitor"],
        },
    })


@app.route("/api/seats/watch", methods=["POST"])
def api_seats_watch():
    body = _body_json()
    user_id = _current_user_id()
    user_store = USER_RUNTIME_STORES.get(user_id)
    result = user_store.add_watcher(
        body.get("course", "机器学习"),
        body.get("student", g.current_user["username"]),
    )
    USER_RUNTIME_STORES.save(user_id)
    return jsonify(result)


@app.route("/api/seats/tick", methods=["POST"])
def api_seats_tick():
    user_id = _current_user_id()
    user_store = USER_RUNTIME_STORES.get(user_id)
    result = user_store.tick_seats()
    result["evidence"] = {
        "mode": "simulation",
        "official_system_connected": False,
        "notice": SIMULATION_FEATURES["seat_monitor"],
    }
    USER_RUNTIME_STORES.save(user_id)
    return jsonify(result)


# ═════════════════════════════════════════════════════════════════════
# API: 教师信息
# ═════════════════════════════════════════════════════════════════════

@app.route("/api/professors")
def api_professors():
    course = _query_arg("course")
    if course:
        return jsonify({"course": course, "teachers": STORE.teachers_for_course(course)})
    return jsonify(STORE.professors_doc)


@app.route("/api/teacher-roster")
def api_teacher_roster():
    rows = STORE.teacher_roster_by_college(
        college=_query_arg("college"),
        q=_query_arg("q"),
        scheduled=_query_arg("scheduled"),
    )
    limit = _query_int("limit", 200)
    teachers = rows if limit < 0 else rows[:limit]
    q = _query_arg("q")
    if q:
        teachers = [STORE.enrich_teacher_row(row) for row in teachers]
    return jsonify({
        "colleges": STORE.teacher_roster_colleges(),
        "teachers": teachers,
        "total": len(rows),
    })


@app.route("/api/faculty-profiles")
def api_faculty_profiles():
    rows = STORE.faculty_profiles(
        college=_query_arg("college"),
        rank=_query_arg("rank"),
        q=_query_arg("q"),
        tutor=_query_arg("tutor"),
    )
    limit = _query_int("limit", 200)
    teachers = rows if limit < 0 else rows[:limit]
    return jsonify({
        "source": STORE.faculty_profiles_doc.get("source", {}),
        "colleges": STORE.faculty_profile_colleges(),
        "ranks": STORE.faculty_profile_ranks(),
        "teachers": teachers,
        "total": len(rows),
    })


# ═════════════════════════════════════════════════════════════════════
# API: 课程难度
# ═════════════════════════════════════════════════════════════════════

@app.route("/api/difficulty")
def api_difficulty():
    course = _query_arg("course")
    if course:
        return jsonify(COURSE_DIFFICULTY.for_course_detail(course))
    return jsonify(COURSE_DIFFICULTY.full_stats())


@app.route("/api/difficulty/search")
def api_difficulty_search():
    q = _query_arg("q")
    top_k = _query_int("top_k", 20)
    return jsonify(COURSE_DIFFICULTY.search(q, top_k) if q else [])


@app.route("/api/difficulty/top")
def api_difficulty_top():
    k = _query_int("k", 20)
    return jsonify(COURSE_DIFFICULTY.top_hardest(k))


# ═════════════════════════════════════════════════════════════════════
# API: LLM / 知识库 状态与开关
# ═════════════════════════════════════════════════════════════════════

@app.route("/api/llm/status")
def api_llm_status():
    return jsonify(_llm_status_for_user())


@app.route("/api/llm/config", methods=["GET", "PUT", "DELETE"])
def api_llm_config():
    """Read or update the current user's encrypted BYOK configuration."""
    user_id = _current_user_id()
    if request.method == "GET":
        config = USER_LLM_CONFIGS.public_settings(user_id)
        config["status"] = _llm_status_for_user()
        return jsonify(config)
    if request.method == "DELETE":
        USER_LLM_CONFIGS.clear_settings(user_id)
        return jsonify({
            **USER_LLM_CONFIGS.public_settings(user_id),
            "status": _llm_status_for_user(),
        })

    try:
        config = USER_LLM_CONFIGS.save_settings(user_id, _body_json())
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    session["llm_enabled"] = True
    return jsonify({
        **config,
        "status": _llm_status_for_user(),
    })


@app.route("/api/llm/test", methods=["POST"])
def api_llm_test():
    """Run a small, user-triggered compatibility check against the active API."""
    user_id = _current_user_id()
    limited = _rate_limit(
        f"llm-test:{user_id}",
        limit=5,
        window_seconds=60,
    )
    if limited:
        return limited
    client = _llm_client_for_user(user_id)
    if not client.api_key:
        return jsonify({
            "error": "当前没有可用的 API Key，请先保存配置",
            "code": "llm_not_configured",
        }), 400
    started = time.perf_counter()
    try:
        result = client.request_chat_completion(
            {
                "model": client.model,
                "messages": [
                    {
                        "role": "user",
                        "content": "只回复：连接成功",
                    }
                ],
            },
            timeout=15,
        )
        content = str(
            result.get("choices", [{}])[0]
            .get("message", {})
            .get("content", "")
        ).strip()
    except (RuntimeError, OSError, ValueError, KeyError, IndexError, json.JSONDecodeError) as exc:
        return jsonify({
            "error": _public_llm_connection_error(exc),
            "code": "llm_test_failed",
            "provider": client.provider,
            "model": client.model,
        }), 502
    return jsonify({
        "status": "ok",
        "message": content or "连接成功",
        "provider": client.provider,
        "model": client.model,
        "latency_ms": round((time.perf_counter() - started) * 1000),
    })


def _public_llm_connection_error(exc: Exception) -> str:
    """Return a safe, actionable message without echoing credentials/provider JSON."""
    message = str(exc).lower()
    if "401" in message or "incorrect api key" in message or "invalid api key" in message:
        return "API Key 无效或已被撤销，请重新创建密钥并保存配置。"
    if "429" in message or "insufficient_quota" in message or "quota" in message:
        return "API 账户额度不足或未开通计费，请充值额度，或切换到其他可用模型。"
    if "404" in message or "model_not_found" in message:
        return "模型名称不存在或当前账号无权使用，请核对模型名称。"
    if "timeout" in message or "timed out" in message:
        return "连接模型服务超时，请检查网络或稍后重试。"
    return "模型服务连接失败，请检查 API 地址、模型名称、密钥和账户状态。"


@app.route("/api/llm/toggle", methods=["POST"])
def api_llm_toggle():
    """切换当前登录用户的大模型启用状态"""
    body = _body_json()
    enabled = body.get("enabled")
    if enabled is not None:
        session["llm_enabled"] = bool(enabled)
    return jsonify(_llm_status_for_user())


@app.route("/api/knowledge/status")
def api_knowledge_status():
    status = _knowledge_base_for().status()
    status.update({
        "total_majors": len(STORE.majors),
        "coverage_complete": status.get("majors_indexed") == len(STORE.majors),
    })
    return jsonify(status)


@app.route("/api/knowledge/toggle", methods=["POST"])
def api_knowledge_toggle():
    """切换知识库（长期记忆）启用状态"""
    body = _body_json()
    enabled = body.get("enabled")
    knowledge_base = _knowledge_base_for()
    if enabled is not None:
        knowledge_base.set_enabled(bool(enabled))
    return jsonify(knowledge_base.status())


@app.route("/api/knowledge/search")
def api_knowledge_search():
    """搜索知识库"""
    q = _query_arg("q", "")
    top_k = _query_int("top_k", 5)
    results = _knowledge_base_for().search(q, top_k=top_k)
    return jsonify({
        "query": q,
        "results": results,
        "total": len(results),
    })


@app.route("/api/knowledge/records")
def api_knowledge_records():
    """列出知识库所有记录"""
    limit = _query_int("limit", 100)
    knowledge_base = _knowledge_base_for()
    return jsonify({
        "records": knowledge_base.all_records(limit=limit),
        "total": knowledge_base.count(),
    })


@app.route("/api/knowledge/clear", methods=["POST"])
def api_knowledge_clear():
    """清空当前用户的对话记忆，保留公共专业知识索引"""
    knowledge_base = _knowledge_base_for()
    count = knowledge_base.clear_conversation_memories()
    return jsonify({"status": "ok", "cleared": count})


# ═════════════════════════════════════════════════════════════════════
# API: 对话管理 (Conversations)
# ═════════════════════════════════════════════════════════════════════

@app.route("/api/conversations")
def api_conversations():
    """列出所有会话"""
    limit = _query_int("limit", 50)
    user_id = _current_user_id()
    conversations = list_conversations(limit=limit, user_id=user_id)
    return jsonify({
        "conversations": conversations,
        "total": len(list_conversations(limit=500, user_id=user_id)),
    })


@app.route("/api/conversations", methods=["POST"])
def api_conversations_create():
    """创建新会话"""
    body = _body_json()
    title = body.get("title", "")
    conv = create_conversation(title=title, user_id=_current_user_id())
    return jsonify(conv), 201


@app.route("/api/conversations/<conv_id>")
def api_conversations_get(conv_id: str):
    """获取单条会话"""
    conv = get_conversation(conv_id, user_id=_current_user_id())
    if not conv:
        return jsonify({"error": "conversation not found"}), 404
    return jsonify(conv)


@app.route("/api/conversations/<conv_id>", methods=["DELETE"])
def api_conversations_delete(conv_id: str):
    """删除会话"""
    user_id = _current_user_id()
    ok = delete_conversation(conv_id, user_id=user_id)
    if not ok:
        return jsonify({"error": "conversation not found"}), 404
    _knowledge_base_for().delete_conversation_memories(conv_id)
    return jsonify({"status": "ok", "deleted": conv_id})


@app.route("/api/conversations/<conv_id>/rename", methods=["POST"])
def api_conversations_rename(conv_id: str):
    """重命名会话"""
    body = _body_json()
    new_title = body.get("title", "").strip()
    if not new_title:
        return jsonify({"error": "title is required"}), 400
    result = rename_conversation(conv_id, new_title, user_id=_current_user_id())
    if not result:
        return jsonify({"error": "conversation not found"}), 404
    return jsonify(result)


@app.route("/api/conversations/<conv_id>/messages", methods=["POST"])
def api_conversations_add_message(conv_id: str):
    """向会话追加消息"""
    body = _body_json()
    role = body.get("role", "user")
    content = body.get("content", "")
    if not content:
        return jsonify({"error": "content is required"}), 400
    result = add_message(conv_id, role, content, user_id=_current_user_id())
    if not result:
        return jsonify({"error": "conversation not found"}), 404
    return jsonify(result)


@app.route("/api/conversations/<conv_id>/messages/<int:msg_index>", methods=["DELETE"])
def api_conversations_delete_message(conv_id: str, msg_index: int):
    """删除会话中的单条消息"""
    result = delete_message(conv_id, msg_index, user_id=_current_user_id())
    if not result:
        return jsonify({"error": "message or conversation not found"}), 404
    return jsonify(result)


# ═════════════════════════════════════════════════════════════════════
# API: 静态数据
# ═════════════════════════════════════════════════════════════════════

@app.route("/api/static-data/<name>")
def api_static_data(name: str):
    safe = Path(name).name
    return jsonify(STORE.__dict__.get(safe, {}))


# ═════════════════════════════════════════════════════════════════════
# API: Function Calling Schema
# ═════════════════════════════════════════════════════════════════════

@app.route("/api/functions")
def api_functions():
    return jsonify({
        "functions": get_function_schemas(),
        "langchain_available": _langchain_available(),
    })


@app.route("/api/functions/execute", methods=["POST"])
def api_functions_execute():
    body = _body_json()
    name = body.get("name", "")
    arguments = body.get("arguments", {})
    if isinstance(arguments, str):
        try:
            arguments = json.loads(arguments)
        except json.JSONDecodeError:
            return jsonify({"error": "Invalid arguments JSON"}), 400

    result = FC_EXECUTOR.execute(name, arguments)
    try:
        return jsonify(json.loads(result))
    except (json.JSONDecodeError, TypeError):
        return jsonify({"result": result})


# ═════════════════════════════════════════════════════════════════════
# API: 聊天端点（增强版 — 集成知识库长期记忆）
# ═════════════════════════════════════════════════════════════════════

@app.route("/api/chat", methods=["POST"])
def api_chat():
    """
    标准聊天端点（非流式）。
    
    新增特性：
    - 支持 conversation_id 实现多会话管理
    - 集成知识库长期记忆（自动检索相关上下文）
    - 自动存储对话到知识库
    """
    limited = _rate_limit(
        f"chat:{_current_user_id()}",
        limit=30,
        window_seconds=60,
    )
    if limited:
        return limited
    body = _body_json()
    message = body.get("message", "")
    context = body.get("context", {})
    conv_id = body.get("conversation_id", "") or context.get("conversation_id", "")
    user_id = _current_user_id()
    knowledge_base = _knowledge_base_for(user_id)

    # 获取或创建活跃会话
    conv = get_or_create_active(conv_id if conv_id else None, user_id=user_id)
    conv_id = conv["id"]

    # 增强上下文：注入知识库长期记忆
    kb_enabled = knowledge_base.enabled and bool(context.get("knowledge_base_enabled", True))
    kb_context = ""
    if kb_enabled and len(message.strip()) > 4:
        kb_context = knowledge_base.get_relevant_context(
            message,
            max_items=5,
            conversation_id=conv_id,
        )

    # 将知识库记忆注入上下文
    enhanced_context = dict(context)
    if kb_context:
        enhanced_context["long_term_memory"] = kb_context
    enhanced_context["conversation_id"] = conv_id
    enhanced_context["user_id"] = user_id
    enhanced_context["knowledge_base_enabled"] = kb_enabled
    enhanced_context["llm_enabled"] = bool(session.get("llm_enabled", True))
    enhanced_context["_runtime_store"] = USER_RUNTIME_STORES.get(user_id)
    enhanced_context["_llm_client"] = _llm_client_for_user(user_id)

    # 调用 Agent
    result = AGENT.respond(message, enhanced_context)
    if result.get("intent") == "seat_watch":
        USER_RUNTIME_STORES.save(user_id)

    # 存储到对话记录
    add_message(conv_id, "user", message, user_id=user_id)
    add_message(
        conv_id,
        "assistant",
        result.get("answer", ""),
        user_id=user_id,
    )

    # 存储到知识库（长期记忆）
    if kb_enabled:
        knowledge_base.store(
            user_message=message,
            assistant_answer=result.get("answer", ""),
            conversation_id=conv_id,
            metadata={
                "intent": result.get("intent", ""),
                "major_id": enhanced_context.get("major_id", ""),
            },
        )

    result["conversation_id"] = conv_id
    return jsonify(result)


@app.route("/api/chat/stream", methods=["POST"])
def api_chat_stream():
    """
    SSE 流式聊天端点。
    支持多会话 + 知识库长期记忆。
    """
    limited = _rate_limit(
        f"chat-stream:{_current_user_id()}",
        limit=30,
        window_seconds=60,
    )
    if limited:
        return limited
    body = _body_json()
    message = body.get("message", "")
    context = body.get("context", {})
    conv_id = body.get("conversation_id", "") or context.get("conversation_id", "")
    user_id = _current_user_id()
    knowledge_base = _knowledge_base_for(user_id)

    conv = get_or_create_active(conv_id if conv_id else None, user_id=user_id)
    conv_id = conv["id"]

    # 知识库上下文
    kb_context = ""
    kb_enabled = knowledge_base.enabled and bool(context.get("knowledge_base_enabled", True))
    if kb_enabled and len(message.strip()) > 4:
        kb_context = knowledge_base.get_relevant_context(
            message,
            max_items=5,
            conversation_id=conv_id,
        )

    enhanced_context = dict(context)
    if kb_context:
        enhanced_context["long_term_memory"] = kb_context
    enhanced_context["conversation_id"] = conv_id
    enhanced_context["user_id"] = user_id
    enhanced_context["knowledge_base_enabled"] = kb_enabled
    enhanced_context["llm_enabled"] = bool(session.get("llm_enabled", True))
    enhanced_context["_runtime_store"] = USER_RUNTIME_STORES.get(user_id)
    enhanced_context["_llm_client"] = _llm_client_for_user(user_id)

    def generate():
        yield _sse("meta", {"status": "started", "message": message, "conversation_id": conv_id})

        add_message(conv_id, "user", message, user_id=user_id)

        llm_status = _llm_status_for_user()
        if llm_status.get("enabled") and _langchain_available():
            yield from _stream_with_agent(
                message,
                enhanced_context,
                llm_status,
                conv_id,
                user_id,
                knowledge_base,
            )
        else:
            yield from _stream_local(
                message,
                enhanced_context,
                conv_id,
                user_id,
                knowledge_base,
            )

        yield _sse("meta", {"status": "done"})

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


def _sse(event: str, data: Any) -> str:
    payload = json.dumps(data, ensure_ascii=False)
    return f"event: {event}\ndata: {payload}\n\n"


def _stream_local(
    message: str,
    context: dict[str, Any],
    conv_id: str,
    user_id: str,
    knowledge_base: KnowledgeBase,
):
    """本地模式流式响应"""
    result = AGENT.respond(message, context)
    if result.get("intent") == "seat_watch":
        USER_RUNTIME_STORES.save(user_id)
    answer = result.get("answer", "")
    intent = result.get("intent", "")
    data = result.get("data", {})

    yield _sse("intent", {"intent": intent})

    import re
    sentences = re.split(r"(?<=[。！？\n])", answer)
    for sentence in sentences:
        if sentence.strip():
            yield _sse("token", {"token": sentence})
            time.sleep(0.05)

    yield _sse("result", {
        **result,
        "intent": intent,
        "data": data,
        "answer": answer,
    })

    # 存储到知识库
    if context.get("knowledge_base_enabled"):
        knowledge_base.store(
            message,
            answer,
            conversation_id=conv_id,
            metadata={"intent": intent},
        )
    add_message(conv_id, "assistant", answer, user_id=user_id)


def _stream_with_agent(
    message: str,
    context: dict[str, Any],
    llm_status: dict[str, Any],
    conv_id: str,
    user_id: str,
    knowledge_base: KnowledgeBase,
):
    """LLM 模式流式响应"""
    llm_client = context.get("_llm_client") or AGENT.llm
    try:
        from langchain.agents import AgentExecutor, create_openai_functions_agent
        from langchain.schema import SystemMessage, HumanMessage, AIMessage
        from langchain_openai import ChatOpenAI
        from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder

        llm_client.validate_endpoint()
        llm = ChatOpenAI(
            model=llm_status.get("model", "gpt-4o-mini"),
            openai_api_key=llm_client.api_key,
            openai_api_base=llm_client.base_url,
            streaming=True,
        )

        tools = _build_langchain_tools()
        if not tools:
            yield _sse("error", {"error": "无法初始化 LangChain 工具"})
            return

        # 构建带有知识库记忆的系统提示
        kb_context = context.get("long_term_memory", "")
        memory_section = (
            f"\n\n【长期记忆参考】\n{kb_context}\n"
            if kb_context else ""
        )
        system_prompt = (
            "你是华侨大学学业规划 AI「勤勉」。你可以使用以下工具来回答用户的问题。\n"
            "必须把工具返回的结构化 JSON 与检索到的知识库上下文结合起来回答，不能只照抄其中一方。\n"
            "工具中的专业、学制、课程、学分、学院和教师等事实优先级最高，所有数值必须保留。\n"
            "当结果包含多门课程、多个学期、多个教师或多个比较项时，必须使用标准 Markdown 表格；"
            "课程规划表至少包含学期、课程、学分、类别和学习重点。\n"
            "职业画像要按学年组织；非毕业学年的第3学期只能作为0学分小学期职业增强建议，"
            "毕业学年只保留两个正式学期。\n"
            "表格后要结合用户问题给出个性化分析，不要输出固定模板式长段落。\n"
            "不要编造未提供的数据；知识库与工具冲突时以结构化工具数据为准，并说明需以学院最新培养方案复核。\n"
            "回答要分层、清晰、适合学生继续追问。\n"
            f"{memory_section}"
        )

        prompt = ChatPromptTemplate.from_messages([
            SystemMessage(content=system_prompt),
            MessagesPlaceholder(variable_name="chat_history", optional=True),
            HumanMessage(content="{input}"),
            MessagesPlaceholder(variable_name="agent_scratchpad"),
        ])

        agent = create_openai_functions_agent(llm, tools, prompt)
        agent_executor = AgentExecutor(
            agent=agent,
            tools=tools,
            verbose=True,
            max_iterations=5,
            handle_parsing_errors=True,
        )

        chat_history = context.get("chat_history", [])
        messages = [SystemMessage(content=system_prompt)]
        for turn in chat_history[-6:]:
            if turn.get("role") == "user":
                messages.append(HumanMessage(content=turn.get("text", "")))
            elif turn.get("role") == "assistant":
                messages.append(AIMessage(content=turn.get("text", "")))

        messages.append(HumanMessage(content=message))

        full_answer = ""
        for chunk in agent_executor.stream({"input": message, "chat_history": chat_history[-6:]}):
            if "output" in chunk:
                token = chunk["output"]
                if token and len(token) > len(full_answer):
                    new_part = token[len(full_answer):]
                    if new_part:
                        yield _sse("token", {"token": new_part})
                    full_answer = token
            elif "intermediate_step" in chunk:
                for step in chunk["intermediate_step"]:
                    if len(step) >= 2:
                        tool_name = step[0].tool if hasattr(step[0], "tool") else ""
                        tool_input = step[0].tool_input if hasattr(step[0], "tool_input") else ""
                        yield _sse("function_call", {
                            "tool": str(tool_name),
                            "input": str(tool_input)[:200],
                        })

        yield _sse("result", {
            "answer": full_answer,
            "answer_mode": "llm_knowledge_hybrid",
            "grounding": {
                "structured_data": True,
                "knowledge_base": bool(kb_context),
                "llm": True,
            },
        })

        if context.get("knowledge_base_enabled"):
            knowledge_base.store(
                message,
                full_answer,
                conversation_id=conv_id,
                metadata={"intent": "agent"},
            )
        add_message(conv_id, "assistant", full_answer, user_id=user_id)

    except ImportError as e:
        llm_client.last_error = str(e)
        yield _sse("error", {
            "error": f"LangChain 未安装或导入失败: {str(e)}",
            "hint": "请安装: pip install langchain langchain-openai",
        })
    except Exception as e:
        llm_client.last_error = str(e)
        yield _sse("error", {"error": f"Agent 执行异常: {str(e)}"})


def _build_langchain_tools() -> list:
    try:
        from langchain.tools import Tool
        return [
            Tool(
                name="analyze_course_hardness",
                func=lambda c="": FC_EXECUTOR.execute("analyze_course_hardness", {"course_name": c}),
                description="分析课程难度 (输入: 课程名称)",
            ),
            Tool(
                name="search_majors",
                func=lambda q="": FC_EXECUTOR.execute("search_majors", {"query": q}),
                description="搜索专业 (输入: 关键词)",
            ),
            Tool(
                name="get_curriculum",
                func=lambda m="": FC_EXECUTOR.execute("get_curriculum", {"major_id": m}),
                description="查询培养方案 (输入: 专业ID)",
            ),
            Tool(
                name="match_professor",
                func=lambda i="": FC_EXECUTOR.execute("match_professor", {"interest": i}),
                description="匹配研究方向对应的老师 (输入: 研究方向)",
            ),
            Tool(
                name="plan_career",
                func=lambda c="": FC_EXECUTOR.execute("plan_career", {"career": c}),
                description="职业规划 (输入: 目标岗位)",
            ),
            Tool(
                name="get_course_teachers",
                func=lambda c="": FC_EXECUTOR.execute("get_course_teachers", {"course_name": c}),
                description="查询课程任课老师 (输入: 课程名称)",
            ),
            Tool(
                name="search_course_difficulty",
                func=lambda k="": FC_EXECUTOR.execute("search_course_difficulty", {"keyword": k}),
                description="按关键词搜索课程难度",
            ),
        ]
    except ImportError:
        return []


# ═════════════════════════════════════════════════════════════════════
# API: 其他 POST 端点
# ═════════════════════════════════════════════════════════════════════

@app.route("/api/plan", methods=["POST"])
def api_plan():
    body = _body_json()
    result = CAREER_PLANNER.plan(
        body.get("career", "算法工程师"),
        body.get("major_id"),
    )
    result["evidence"] = {
        "mode": "reference_inference",
        "official_employment_outcome": False,
        "notice": SIMULATION_FEATURES["career_plan"],
    }
    return jsonify(result)


@app.route("/api/exports", methods=["POST"])
def api_exports():
    limited = _rate_limit(
        f"exports:{_current_user_id()}",
        limit=15,
        window_seconds=60,
    )
    if limited:
        return limited
    body = _body_json()
    if len(json.dumps(body, ensure_ascii=False)) > 2_000_000:
        return jsonify({"error": "导出数据过大，请缩小导出范围"}), 413
    try:
        output, filename, mimetype = build_export(
            body.get("kind", ""),
            body.get("format", ""),
            body.get("title", ""),
            body.get("data") or {},
        )
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    except ImportError as exc:
        return jsonify({"error": f"服务器缺少导出组件：{exc.name}"}), 503
    return send_file(
        output,
        mimetype=mimetype,
        as_attachment=True,
        download_name=filename,
        max_age=0,
    )


@app.route("/api/course/analyze", methods=["POST"])
def api_course_analyze():
    body = _body_json()
    return jsonify(
        COURSE_ANALYZER.analyze(body.get("course", "数据结构"), body.get("reviews") or [])
    )


@app.route("/api/conflicts", methods=["POST"])
def api_conflicts():
    body = _body_json()
    result = CONFLICT_RESOLVER.resolve(
        body.get("major_id", ""),
        body.get("selected_courses", []),
    )
    result["evidence"] = {
        "mode": "decision_support",
        "official_system_connected": False,
        "notice": SIMULATION_FEATURES["conflict_resolution"],
    }
    return jsonify(result)


@app.route("/api/professors/match", methods=["POST"])
def api_professors_match():
    body = _body_json()
    return jsonify(
        PROF_MATCHER.match(body.get("interest_text", ""), int(body.get("top_k", 5)))
    )


@app.route("/api/credits/check", methods=["POST"])
def api_credits_check():
    body = _body_json()
    return jsonify(
        CREDIT_CHECKER.check(
            body.get("major_id", ""),
            body.get("completed_courses", []),
            body.get("student_type", "domestic"),
        )
    )


@app.route("/api/import/teacher-schedule", methods=["POST"])
def api_import_teacher_schedule():
    body = _body_json()
    return jsonify(
        STORE.import_teacher_schedule_text(body.get("text", ""), bool(body.get("replace", False)))
    )


@app.route("/api/import/demo", methods=["POST"])
def api_import_demo():
    body = _body_json()
    imported = {}
    if "reviews" in body:
        STORE.reviews_doc["reviews"].extend(body["reviews"])
        imported["reviews"] = len(body["reviews"])
    if "professors" in body:
        STORE.professors_doc["professors"].extend(body["professors"])
        imported["professors"] = len(body["professors"])
    if "offerings" in body:
        STORE.seat_doc["offerings"].extend(body["offerings"])
        imported["offerings"] = len(body["offerings"])
    return jsonify({"status": "ok", "imported": imported})


# ═════════════════════════════════════════════════════════════════════
# 错误处理
# ═════════════════════════════════════════════════════════════════════

@app.errorhandler(404)
def not_found(e):
    return jsonify({"error": "not found"}), 404


@app.errorhandler(500)
def server_error(e):
    return jsonify({"error": str(e)}), 500


# ═════════════════════════════════════════════════════════════════════
# API: 图片/文件上传分析
# ═════════════════════════════════════════════════════════════════════

MAX_UPLOAD_BYTES = 15 * 1024 * 1024
TIMETABLE_PROMPT = (
    "你是华侨大学课表与学业文件分析助手。请读取用户上传的内容，提取所有能确认的课程信息。\n"
    "要求：\n"
    "1. 课程字段包含 name、day、start、end、category、semester。\n"
    "2. day 统一为周一到周日，时间统一为 HH:MM；无法确认的字段留空，不要编造。\n"
    "3. intent 只能是 conflict、career_plan、curriculum、credit_check、general 之一。\n"
    "4. summary 简要说明识别结果和不确定项。\n"
    "5. 只返回一个合法 JSON 对象，不要使用 Markdown 代码块。\n"
    '返回格式：{"intent":"conflict","courses":[{"name":"课程名","day":"周一",'
    '"start":"08:00","end":"09:40","category":"专业必修","semester":1}],'
    '"summary":"...","career":"","major":""}'
)


def _parse_llm_json(answer: str) -> dict[str, Any]:
    cleaned = answer.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", cleaned, flags=re.IGNORECASE)
    try:
        parsed = json.loads(cleaned)
        if isinstance(parsed, dict):
            return parsed
    except json.JSONDecodeError:
        pass

    decoder = json.JSONDecoder()
    for match in re.finditer(r"\{", cleaned):
        try:
            parsed, _ = decoder.raw_decode(cleaned[match.start():])
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            continue
    return {"intent": "general", "summary": cleaned, "courses": []}


def _normalize_analysis_result(result: dict[str, Any]) -> dict[str, Any]:
    normalized_courses = []
    day_aliases = {
        "星期一": "周一", "星期二": "周二", "星期三": "周三", "星期四": "周四",
        "星期五": "周五", "星期六": "周六", "星期日": "周日", "星期天": "周日",
        "Monday": "周一", "Tuesday": "周二", "Wednesday": "周三", "Thursday": "周四",
        "Friday": "周五", "Saturday": "周六", "Sunday": "周日",
    }
    for course in result.get("courses", []) if isinstance(result.get("courses"), list) else []:
        if not isinstance(course, dict) or not str(course.get("name", "")).strip():
            continue
        day = str(course.get("day", "")).strip()
        normalized_courses.append({
            "name": str(course.get("name", "")).strip(),
            "day": day_aliases.get(day, day),
            "start": str(course.get("start", "")).strip(),
            "end": str(course.get("end", "")).strip(),
            "category": str(course.get("category", "")).strip() or "未分类",
            "semester": course.get("semester", ""),
        })
    intent = result.get("intent", "general")
    if intent not in {"conflict", "career_plan", "curriculum", "credit_check", "general"}:
        intent = "general"
    return {
        "intent": intent,
        "courses": normalized_courses,
        "summary": str(result.get("summary", "")).strip(),
        "career": str(result.get("career", "")).strip(),
        "major": str(result.get("major", "")).strip(),
    }


def _call_file_analysis(
    content: Any,
    prompt: str = "",
    *,
    llm: Any | None = None,
    use_vision_model: bool = False,
) -> dict[str, Any]:
    llm = llm or _llm_client_for_user()
    if not llm.api_key:
        raise RuntimeError("未配置大模型 API Key")
    payload = {
        "model": (llm.vision_model or llm.model) if use_vision_model else llm.model,
        "messages": [{"role": "user", "content": content}],
        "max_tokens": 3000,
    }
    try:
        response_data = llm.request_chat_completion(
            payload,
            timeout=max(llm.timeout, 60),
        )
    except RuntimeError as exc:
        raise RuntimeError(f"模型接口调用失败：{str(exc)[:300]}") from exc
    answer = response_data["choices"][0]["message"]["content"]
    return _normalize_analysis_result(_parse_llm_json(answer))


def _llm_likely_supports_images(llm: Any) -> bool:
    """Best-effort capability check for OpenAI-compatible chat providers."""
    provider = str(getattr(llm, "provider", "")).strip().lower()
    base_url = str(getattr(llm, "base_url", "")).strip().lower()
    model = str(
        getattr(llm, "vision_model", "") or getattr(llm, "model", "")
    ).strip().lower()
    vision_markers = (
        "vision", "gpt-4o", "gpt-4.1", "gpt-5", "qwen-vl", "qwen2-vl",
        "qwen2.5-vl", "qvq", "gemini", "claude-3", "claude-sonnet",
        "pixtral", "llava", "internvl", "glm-4v", "doubao-vision",
    )
    if any(marker in model for marker in vision_markers):
        return True
    if provider == "deepseek" or "api.deepseek.com" in base_url:
        return False
    if provider in {"qwen", "dashscope", "tongyi"}:
        return False
    # OpenAI and unknown compatible gateways may support multimodal messages.
    # If they reject image_url, _analyze_image_bytes falls back to local OCR.
    return True


def _is_image_content_unsupported(exc: Exception) -> bool:
    message = str(exc).lower()
    markers = (
        "unknown variant `image_url`",
        "unknown variant 'image_url'",
        'unknown variant "image_url"',
        "does not support image",
        "image input is not supported",
        "image_url is not supported",
        "expected `text`",
        "expected 'text'",
        'expected "text"',
        "multimodal is not supported",
    )
    return any(marker in message for marker in markers)


def _friendly_model_failure(exc: Exception) -> str:
    message = str(exc).lower()
    if "401" in message or "incorrect api key" in message or "invalid api key" in message:
        return "当前大模型 API Key 无效，已改用服务器 OCR 与本地课程规则识别。"
    if "429" in message or "insufficient_quota" in message or "quota" in message:
        return "当前大模型额度不足，已改用服务器 OCR 与本地课程规则识别。"
    if _is_image_content_unsupported(exc):
        return "当前模型不支持直接识图，已改用服务器 OCR 兼容识别。"
    return "大模型暂时不可用，已改用服务器 OCR 与本地课程规则识别。"


def _extract_image_text(data: bytes, filename: str) -> str:
    """Run local Tesseract OCR and preserve word coordinates for timetable grids."""
    suffix = Path(filename).suffix.lower()
    if suffix not in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}:
        suffix = ".png"
    try:
        with tempfile.TemporaryDirectory(prefix="qinmian-ocr-") as temp_dir:
            image_path = Path(temp_dir) / f"upload{suffix}"
            image_path.write_bytes(data)
            completed = subprocess.run(
                [
                    "tesseract",
                    str(image_path),
                    "stdout",
                    "-l",
                    "chi_sim+eng",
                    "--psm",
                    "6",
                    "tsv",
                ],
                check=True,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=45,
            )
    except FileNotFoundError as exc:
        raise RuntimeError("服务器 OCR 组件尚未安装") from exc
    except subprocess.TimeoutExpired as exc:
        raise RuntimeError("图片文字识别超时，请裁剪图片后重试") from exc
    except subprocess.CalledProcessError as exc:
        detail = (exc.stderr or "").strip()
        raise RuntimeError(f"图片文字识别失败：{detail[:160] or '无法读取图片'}") from exc

    lines: dict[tuple[int, int, int, int], list[tuple[int, int, str]]] = {}
    for raw_line in completed.stdout.splitlines()[1:]:
        columns = raw_line.split("\t", 11)
        if len(columns) != 12:
            continue
        text = columns[11].strip()
        if not text:
            continue
        try:
            key = tuple(int(columns[index]) for index in (1, 2, 3, 4))
            left = int(columns[6])
            top = int(columns[7])
        except ValueError:
            continue
        lines.setdefault(key, []).append((top, left, text))

    layout_lines = []
    for words in lines.values():
        words.sort(key=lambda item: item[1])
        top = min(item[0] for item in words)
        positioned = " ".join(f"[x={left}]{text}" for _, left, text in words)
        layout_lines.append((top, positioned))
    layout_lines.sort(key=lambda item: item[0])
    return "\n".join(f"[y={top}] {line}" for top, line in layout_lines).strip()


def _parse_ocr_layout_courses(ocr_text: str) -> list[dict[str, Any]]:
    """Best-effort timetable parser used when the configured LLM is unavailable."""
    positioned_lines: list[dict[str, Any]] = []
    token_pattern = re.compile(r"\[x=(\d+)\]([^\[]+)")
    for raw_line in str(ocr_text or "").splitlines():
        y_match = re.match(r"\[y=(\d+)\]\s*(.*)", raw_line.strip())
        if not y_match:
            continue
        y = int(y_match.group(1))
        tokens = [
            (int(match.group(1)), match.group(2).strip())
            for match in token_pattern.finditer(y_match.group(2))
            if match.group(2).strip()
        ]
        if tokens:
            positioned_lines.append({"y": y, "tokens": tokens})

    day_aliases = {
        "周一": "周一", "星期一": "周一", "礼拜一": "周一",
        "周二": "周二", "星期二": "周二", "礼拜二": "周二",
        "周三": "周三", "星期三": "周三", "礼拜三": "周三",
        "周四": "周四", "星期四": "周四", "礼拜四": "周四",
        "周五": "周五", "星期五": "周五", "礼拜五": "周五",
        "周六": "周六", "星期六": "周六", "礼拜六": "周六",
        "周日": "周日", "星期日": "周日", "星期天": "周日", "周天": "周日",
    }
    day_columns: list[tuple[int, str]] = []
    for line in positioned_lines:
        for x, text in line["tokens"]:
            compact = re.sub(r"\s+", "", text)
            for alias, normalized in day_aliases.items():
                if alias in compact:
                    day_columns.append((x, normalized))
                    break
    # One heading can be recognized twice; retain one x position per weekday.
    unique_days: dict[str, int] = {}
    for x, day in day_columns:
        unique_days.setdefault(day, x)
    day_columns = sorted((x, day) for day, x in unique_days.items())

    time_pattern = re.compile(r"(?<!\d)([01]?\d|2[0-3])[:：.]([0-5]\d)(?!\d)")
    time_rows: list[tuple[int, str, str]] = []
    for line in positioned_lines:
        joined = " ".join(text for _, text in line["tokens"])
        times = [
            f"{int(hour):02d}:{minute}"
            for hour, minute in time_pattern.findall(joined)
        ]
        if not times:
            continue
        start = times[0]
        if len(times) >= 2:
            end = times[1]
        else:
            hour, minute = map(int, start.split(":"))
            total = hour * 60 + minute + 100
            end = f"{(total // 60) % 24:02d}:{total % 60:02d}"
        time_rows.append((line["y"], start, end))

    def normalized(value: str) -> str:
        return re.sub(r"[^0-9a-z\u4e00-\u9fff]+", "", value.lower())

    known_courses = _known_ocr_course_names()
    results: list[dict[str, Any]] = []
    seen: set[tuple[str, str, str]] = set()
    for line in positioned_lines:
        token_norms = [normalized(text) for _, text in line["tokens"]]
        joined_norm = "".join(token_norms)
        if not joined_norm:
            continue
        occupied: list[tuple[int, int]] = []
        for course_name in known_courses:
            course_norm = normalized(course_name)
            character_index = joined_norm.find(course_norm)
            if character_index < 0:
                continue
            span = (character_index, character_index + len(course_norm))
            if any(span[0] < used[1] and used[0] < span[1] for used in occupied):
                continue
            running = 0
            course_x = line["tokens"][0][0]
            for (x, _), token_norm in zip(line["tokens"], token_norms):
                if running + len(token_norm) > character_index:
                    course_x = x
                    break
                running += len(token_norm)
            day = ""
            if day_columns:
                day = min(day_columns, key=lambda item: abs(item[0] - course_x))[1]
            if not day:
                raw_text = "".join(text for _, text in line["tokens"])
                for alias, normalized_day in day_aliases.items():
                    if alias in raw_text:
                        day = normalized_day
                        break
            start = end = ""
            if time_rows:
                _, start, end = min(
                    time_rows,
                    key=lambda item: abs(item[0] - line["y"]),
                )
            key = (course_name, day, start)
            if key in seen:
                continue
            seen.add(key)
            occupied.append(span)
            results.append({
                "name": course_name,
                "day": day,
                "start": start,
                "end": end,
                "category": "未分类",
                "semester": "",
            })
    return results


def _local_ocr_analysis(ocr_text: str, warning: str) -> dict[str, Any]:
    courses = _parse_ocr_layout_courses(ocr_text)
    if courses:
        summary = (
            f"大模型当前不可用，服务器已通过 OCR 和本地课程库识别出 "
            f"{len(courses)} 门课程。请核对星期和时间后再执行冲突检查。"
        )
    else:
        summary = (
            "服务器已经读取图片文字，但没有可靠匹配到课程。"
            "请上传更清晰、仅包含课表区域的截图，或手动添加课程。"
        )
    return {
        "intent": "conflict",
        "courses": courses,
        "summary": summary,
        "career": "",
        "major": "",
        "analysis_method": "local_ocr_rules",
        "analysis_note": warning,
    }


@functools.lru_cache(maxsize=1)
def _known_ocr_course_names() -> tuple[str, ...]:
    def normalized(value: str) -> str:
        return re.sub(r"[^0-9a-z\u4e00-\u9fff]+", "", value.lower())

    return tuple(sorted(
        {
            name.strip()
            for name in STORE.all_course_names()
            if isinstance(name, str) and len(normalized(name)) >= 2
        },
        key=lambda name: len(normalized(name)),
        reverse=True,
    ))


def _decode_text_file(data: bytes) -> str:
    for encoding in ("utf-8-sig", "gb18030", "utf-16"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_document_text(data: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md", ".csv", ".json"}:
        return _decode_text_file(data)
    if suffix in {".xlsx", ".xlsm"}:
        from openpyxl import load_workbook
        workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        lines = []
        for sheet in workbook.worksheets[:8]:
            lines.append(f"工作表：{sheet.title}")
            for row in sheet.iter_rows(max_row=300, values_only=True):
                values = [str(value).strip() for value in row if value not in (None, "")]
                if values:
                    lines.append("\t".join(values))
        return "\n".join(lines)
    if suffix == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages[:30])
    if suffix == ".docx":
        from docx import Document
        document = Document(io.BytesIO(data))
        lines = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                lines.append("\t".join(cell.text.strip() for cell in row.cells))
        return "\n".join(lines)
    raise ValueError("暂不支持该文件格式。支持图片、PDF、DOCX、XLSX、CSV、TXT、JSON。")


def _analyze_image_bytes(data: bytes, filename: str, mime_type: str, prompt: str = "") -> dict[str, Any]:
    llm = _llm_client_for_user()
    model_error: RuntimeError | None = None
    analysis_prompt = TIMETABLE_PROMPT
    if prompt:
        analysis_prompt += f"\n\n用户补充要求：{prompt}"

    if _llm_likely_supports_images(llm):
        encoded = base64.b64encode(data).decode("ascii")
        content = [
            {"type": "text", "text": analysis_prompt},
            {
                "type": "image_url",
                "image_url": {
                    "url": f"data:{mime_type};base64,{encoded}",
                    "detail": "high",
                },
            },
        ]
        try:
            result = _call_file_analysis(
                content,
                prompt,
                llm=llm,
                use_vision_model=True,
            )
            result["analysis_method"] = "vision_model"
            return result
        except RuntimeError as exc:
            model_error = exc

    ocr_text = _extract_image_text(data, filename)
    if not ocr_text:
        raise RuntimeError(
            "当前模型不支持直接识图，服务器也没有从图片中识别到文字。"
            "请上传更清晰、方向正确的课表截图，或改用支持视觉输入的模型。"
        )
    text_content = (
        f"{analysis_prompt}\n\n"
        "下面是服务器从图片中提取的 OCR 文字。x、y 是文字在图片中的位置，"
        "请结合坐标还原课表的行列关系；不确定的内容不要编造。\n"
        f"文件名：{filename}\nOCR 布局文字：\n{ocr_text[:60000]}"
    )
    # If a direct visual request already proved that authentication/quota is
    # unavailable, do not send the same failing credential a second time.
    if model_error and not _is_image_content_unsupported(model_error):
        return _local_ocr_analysis(ocr_text, _friendly_model_failure(model_error))
    try:
        result = _call_file_analysis(text_content, prompt, llm=llm)
        result["analysis_method"] = "local_ocr_text_model"
        result["analysis_note"] = (
            "当前模型不支持直接识图，已自动使用服务器 OCR 兼容识别。"
        )
        return result
    except RuntimeError as exc:
        return _local_ocr_analysis(ocr_text, _friendly_model_failure(exc))


@app.route("/api/files/analyze", methods=["POST"])
def api_file_analyze():
    limited = _rate_limit(
        f"file-analysis:{_current_user_id()}",
        limit=10,
        window_seconds=10 * 60,
    )
    if limited:
        return limited
    uploaded = request.files.get("file")
    if not uploaded or not uploaded.filename:
        return jsonify({"error": "请选择要上传的文件"}), 400
    data = uploaded.read(MAX_UPLOAD_BYTES + 1)
    if len(data) > MAX_UPLOAD_BYTES:
        return jsonify({"error": "文件不能超过 15 MB"}), 413
    prompt = request.form.get("prompt", "").strip()
    mime_type = uploaded.mimetype or mimetypes.guess_type(uploaded.filename)[0] or "application/octet-stream"
    try:
        if mime_type.startswith("image/"):
            result = _analyze_image_bytes(data, uploaded.filename, mime_type, prompt)
        else:
            document_text = _extract_document_text(data, uploaded.filename).strip()
            if not document_text:
                return jsonify({"error": "文件中没有可读取的文字，请上传清晰截图或可复制文本的文件"}), 400
            analysis_prompt = TIMETABLE_PROMPT
            if prompt:
                analysis_prompt += f"\n\n用户补充要求：{prompt}"
            content = f"{analysis_prompt}\n\n文件名：{uploaded.filename}\n文件内容：\n{document_text[:60000]}"
            result = _call_file_analysis(content, prompt)
        result["filename"] = uploaded.filename
        result["file_type"] = mime_type
        return jsonify(result)
    except (ValueError, ImportError) as exc:
        return jsonify({"error": str(exc)}), 400
    except Exception as exc:
        return jsonify({"error": f"文件识别失败：{str(exc)[:300]}"}), 502

@app.route("/api/vision/analyze", methods=["POST"])
def api_vision_analyze():
    """
    接收上传的图片（课表截图等）并通过 LLM 分析。
    
    请求体：
    {
        "image": "base64编码的图片数据",
        "filename": "文件名",
        "prompt": "可选，自定义分析提示"
    }
    
    返回：
    {
        "intent": "conflict|career_plan|curriculum|credit_check|general",
        "courses": [{"name": "...", "day": "...", "start": "...", "end": "...", "category": "..."}],
        "summary": "分析总结文字",
        "career": "识别的岗位（如果是职业规划）",
        "major": "识别的专业"
    }
    """
    body = _body_json()
    image_b64 = body.get("image", "")
    filename = body.get("filename", "image.jpg")
    user_prompt = body.get("prompt", "")
    
    if not image_b64:
        return jsonify({"error": "no image data"}), 400
    
    try:
        if image_b64.startswith("data:"):
            header, image_b64 = image_b64.split(",", 1)
            mime_type = header.split(";")[0].replace("data:", "")
        else:
            mime_type = mimetypes.guess_type(filename)[0] or "image/jpeg"
        image_data = base64.b64decode(image_b64, validate=True)
        if len(image_data) > MAX_UPLOAD_BYTES:
            return jsonify({"error": "图片不能超过 15 MB"}), 413
        return jsonify(_analyze_image_bytes(image_data, filename, mime_type, user_prompt))
    except (binascii.Error, ValueError):
        return jsonify({"error": "图片数据格式无效"}), 400
    except Exception as exc:
        return jsonify({"error": f"图片识别失败：{str(exc)[:300]}"}), 502


def main() -> None:
    port = int(os.getenv("PORT", "8765"))
    debug = os.getenv("FLASK_DEBUG", "0") == "1"
    print(f"[勤勉 v4] Flask server starting on http://127.0.0.1:{port}")
    print(f"[勤勉 v4] SSE streaming: http://127.0.0.1:{port}/api/chat/stream")
    print(f"[勤勉 v4] Conversations: http://127.0.0.1:{port}/api/conversations")
    print(f"[勤勉 v4] Knowledge Base: http://127.0.0.1:{port}/api/knowledge/status")
    app.run(host=os.getenv("HOST", "0.0.0.0"), port=port, debug=debug, threaded=True)


if __name__ == "__main__":
    main()
