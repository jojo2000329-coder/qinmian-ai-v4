from __future__ import annotations

import html
import io
import re
from datetime import datetime
from typing import Any


EXPORT_FORMATS = {"docx", "pdf", "xls"}
EXPORT_KINDS = {"conversation", "career_plan"}


def _clean_text(value: Any) -> str:
    text = html.unescape(str(value or ""))
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"```(?:\w+)?\s*", "", text)
    text = text.replace("```", "")
    return text.strip()


def _safe_filename(value: Any, fallback: str = "勤勉导出") -> str:
    text = re.sub(r'[\\/:*?"<>|\x00-\x1f]+', "_", _clean_text(value))
    text = re.sub(r"\s+", " ", text).strip(" ._")
    return (text[:80] or fallback)


def _conversation_rows(payload: dict[str, Any]) -> list[list[str]]:
    rows = [["角色", "内容"]]
    for item in payload.get("messages", [])[:500]:
        if not isinstance(item, dict):
            continue
        role = "用户" if item.get("role") == "user" else "勤勉 AI"
        content = _clean_text(item.get("content", item.get("text", "")))
        if content:
            rows.append([role, content])
    return rows


def _career_rows(payload: dict[str, Any]) -> list[list[Any]]:
    rows: list[list[Any]] = [
        ["目标岗位", "当前专业", "匹配度", "匹配结论", "学年", "学期/阶段", "课程/建议", "类别", "学分", "来源", "说明"]
    ]
    role = _clean_text(payload.get("matched_role") or payload.get("career"))
    major = payload.get("selected_major") or {}
    major_name = _clean_text(major.get("display_name") or major.get("name"))
    fit = payload.get("selected_major_fit") or {}
    periods = payload.get("planning_periods") or payload.get("semesters") or []
    for period in periods[:24]:
        if not isinstance(period, dict):
            continue
        year_label = _clean_text(period.get("year_label"))
        period_label = _clean_text(period.get("short_label") or period.get("label"))
        for course in period.get("courses", [])[:100]:
            if not isinstance(course, dict):
                continue
            rows.append(
                [
                    role,
                    major_name,
                    f"{fit.get('score', 0)}%",
                    _clean_text(fit.get("level")),
                    year_label,
                    period_label,
                    _clean_text(course.get("name")),
                    _clean_text(course.get("category")),
                    course.get("credits", 0),
                    (
                        "小学期职业规划建议"
                        if period.get("term_type") == "summer"
                        else "规划建议"
                        if course.get("origin") == "career"
                        else "培养方案"
                    ),
                    _clean_text(course.get("planning_note")),
                ]
            )
    return rows


def _rows_for(kind: str, payload: dict[str, Any]) -> list[list[Any]]:
    if kind == "conversation":
        return _conversation_rows(payload)
    return _career_rows(payload)


def _build_docx(kind: str, title: str, payload: dict[str, Any]) -> io.BytesIO:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Pt

    document = Document()
    document.add_heading(title, level=0)
    document.add_paragraph(f"导出时间：{datetime.now():%Y-%m-%d %H:%M:%S}")
    if kind == "conversation":
        for role, content in _conversation_rows(payload)[1:]:
            document.add_heading(role, level=1)
            for paragraph in content.splitlines() or [""]:
                document.add_paragraph(paragraph)
    else:
        rows = _career_rows(payload)
        table = document.add_table(rows=1, cols=len(rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for index, value in enumerate(rows[0]):
            table.rows[0].cells[index].text = str(value)
        for row in rows[1:]:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = str(value)
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10.5)
    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output


def _pdf_paragraph(value: Any, style, font_name: str):
    from reportlab.platypus import Paragraph

    safe = html.escape(_clean_text(value)).replace("\n", "<br/>")
    return Paragraph(safe or " ", style.clone(f"cell-{id(value)}", fontName=font_name))


def _build_pdf(kind: str, title: str, payload: dict[str, Any]) -> io.BytesIO:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import PageBreak, SimpleDocTemplate, Spacer, Table, TableStyle

    font_name = "STSong-Light"
    pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    output = io.BytesIO()
    page_size = A4 if kind == "conversation" else landscape(A4)
    document = SimpleDocTemplate(
        output,
        pagesize=page_size,
        leftMargin=24,
        rightMargin=24,
        topMargin=28,
        bottomMargin=28,
        title=title,
    )
    styles = getSampleStyleSheet()
    title_style = styles["Title"].clone("ChineseTitle", fontName=font_name, alignment=TA_CENTER)
    heading_style = styles["Heading2"].clone("ChineseHeading", fontName=font_name)
    body_style = styles["BodyText"].clone(
        "ChineseBody", fontName=font_name, fontSize=9.5, leading=14
    )
    story = [_pdf_paragraph(title, title_style, font_name), Spacer(1, 10)]
    story.append(_pdf_paragraph(f"导出时间：{datetime.now():%Y-%m-%d %H:%M:%S}", body_style, font_name))
    story.append(Spacer(1, 12))
    if kind == "conversation":
        for index, (role, content) in enumerate(_conversation_rows(payload)[1:]):
            story.append(_pdf_paragraph(role, heading_style, font_name))
            story.append(_pdf_paragraph(content, body_style, font_name))
            story.append(Spacer(1, 8))
            if index and index % 12 == 0:
                story.append(PageBreak())
    else:
        rows = _career_rows(payload)
        table_data = [
            [_pdf_paragraph(value, body_style, font_name) for value in row]
            for row in rows
        ]
        widths = [48, 66, 38, 44, 32, 54, 84, 46, 30, 58, 106]
        table = Table(table_data, colWidths=widths, repeatRows=1, hAlign="CENTER")
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbeafe")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#94a3b8")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                    ("LEFTPADDING", (0, 0), (-1, -1), 3),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(table)
    document.build(story)
    output.seek(0)
    return output


def _build_xls(kind: str, title: str, payload: dict[str, Any]) -> io.BytesIO:
    import xlwt

    workbook = xlwt.Workbook(encoding="utf-8")
    sheet = workbook.add_sheet("对话" if kind == "conversation" else "课表")
    header_style = xlwt.easyxf(
        "font: bold on, colour dark_blue; pattern: pattern solid, fore_colour ice_blue;"
        "align: vert centre, horiz centre; borders: left thin, right thin, top thin, bottom thin;"
    )
    body_style = xlwt.easyxf(
        "align: vert top, wrap on; borders: left thin, right thin, top thin, bottom thin;"
    )
    rows = _rows_for(kind, payload)
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            text = str(value)[:32767]
            sheet.write(
                row_index,
                column_index,
                text,
                header_style if row_index == 0 else body_style,
            )
    for column_index in range(len(rows[0])):
        max_length = max(
            [len(str(row[column_index])) for row in rows[:200] if column_index < len(row)] or [8]
        )
        sheet.col(column_index).width = min(max(max_length + 2, 10), 40) * 256
    sheet.freeze_panes = True
    sheet.horz_split_pos = 1
    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    return output


def build_export(
    kind: str,
    export_format: str,
    title: str,
    payload: dict[str, Any],
) -> tuple[io.BytesIO, str, str]:
    kind = str(kind or "").strip().lower()
    export_format = str(export_format or "").strip().lower()
    if kind not in EXPORT_KINDS:
        raise ValueError("不支持的导出数据类型")
    if export_format not in EXPORT_FORMATS:
        raise ValueError("不支持的导出格式")
    if not isinstance(payload, dict):
        raise ValueError("导出数据格式无效")
    rows = _rows_for(kind, payload)
    if len(rows) <= 1:
        raise ValueError("没有可导出的数据")
    safe_title = _safe_filename(title, "勤勉对话" if kind == "conversation" else "勤勉课表")
    builders = {
        "docx": _build_docx,
        "pdf": _build_pdf,
        "xls": _build_xls,
    }
    mime_types = {
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf": "application/pdf",
        "xls": "application/vnd.ms-excel",
    }
    output = builders[export_format](kind, safe_title, payload)
    suffix = "对话" if kind == "conversation" else "职业课表"
    return output, f"{safe_title}_{suffix}.{export_format}", mime_types[export_format]
